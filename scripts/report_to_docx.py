#!/usr/bin/env python3
import json
import sys
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_TEMPLATE = ROOT / "templates" / "weekly-report-template.docx"


def clear_body(document):
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def style_exists(document, style_name):
    try:
        document.styles[style_name]
        return True
    except KeyError:
        return False


def add_paragraph(document, text, style_name="Normal"):
    paragraph = document.add_paragraph(style=style_name if style_exists(document, style_name) else "Normal")
    paragraph.add_run(text)
    return paragraph


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def apply_table_text(cell, text, bold=False):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.08
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(30, 41, 51)


def add_overview_table(document, rows):
    if not rows:
        return
    table = document.add_table(rows=0, cols=3)
    table.style = "Normal Table"
    set_table_width(table)
    widths = [2400, 2200, 4760]
    for row_index, row_values in enumerate(rows):
        cells = table.add_row().cells
        for index, cell in enumerate(cells):
            set_cell_width(cell, widths[index])
            if row_index == 0:
                set_cell_shading(cell, "F2F4F7")
            apply_table_text(cell, row_values[index] if index < len(row_values) else "", bold=row_index == 0)
    document.add_paragraph()


def build_docx(payload, output):
    template = Path(payload.get("templatePath") or DEFAULT_TEMPLATE)
    document = Document(template) if template.exists() else Document()
    clear_body(document)

    model = payload.get("model")
    if model:
        title_text = model.get("title") or "项目周报"
        add_paragraph(document, title_text, "Title")
        add_paragraph(document, model.get("subtitle") or "", "Subtitle")
        if model.get("callout"):
            add_paragraph(document, model["callout"], "Callout")
        if len(model.get("overviewRows") or []) > 1:
            add_paragraph(document, "项目状态总览", "Heading 1")
            add_overview_table(document, model["overviewRows"])
        for section in model.get("sections", []):
            add_paragraph(document, section.get("heading") or "", "Heading 1")
            for paragraph in section.get("paragraphs", []):
                add_paragraph(document, paragraph, "Normal")
    else:
        report = payload["report"]
        lines = [line.strip() for line in report.splitlines() if line.strip()]
        title_text = lines[0] if lines else "项目周报"
        add_paragraph(document, title_text, "Title")
        start = payload["range"]["start"].replace("-", ".")
        end = payload["range"]["end"].replace("-", ".")
        source = payload.get("sourceName") or "生成资料"
        add_paragraph(document, f"汇报周期：{start}-{end}｜来源：{source}", "Subtitle")
        for line in lines[1:]:
            if line.startswith("管理判断："):
                add_paragraph(document, line, "Callout")
            else:
                add_paragraph(document, line, "Normal")

    document.save(output)

    with ZipFile(output) as docx:
        xml = docx.read("word/document.xml").decode("utf-8", errors="ignore")
        if title_text not in xml and "项目周报" not in xml:
            raise RuntimeError("DOCX content verification failed.")


def main():
    if len(sys.argv) != 3:
        raise SystemExit("Usage: report_to_docx.py <payload.json> <output.docx>")
    payload = json.loads(Path(sys.argv[1]).read_text(encoding="utf-8"))
    build_docx(payload, sys.argv[2])


if __name__ == "__main__":
    main()
